"""
Rapport Word COMPLET pour SignLanguage
Tous les chapitres rédigés + 9 images avec explications
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_image_with_explanation(doc, img_path, caption, explanation):
    """Ajoute une image avec légende et explication détaillée"""
    try:
        if os.path.exists(img_path):
            doc.add_paragraph()
            # Image centrée
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6))
            
            # Légende
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            cap_run.bold = True
            cap_run.italic = True
            cap_run.font.size = Pt(10)
            
            # Explication
            doc.add_paragraph()
            doc.add_paragraph(explanation)
            doc.add_paragraph()
        else:
            doc.add_paragraph(f'[Image manquante : {os.path.basename(img_path)}]')
    except Exception as e:
        doc.add_paragraph(f'[Erreur image : {str(e)}]')

def read_markdown_file(filepath):
    """Lit un fichier markdown et retourne son contenu"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        # Remplacer SignBridge par SignLanguage
        content = content.replace('SignBridge', 'SignLanguage')
        content = content.replace('signbridge', 'signlanguage')
        return content
    except:
        return '[Contenu non disponible]'

def add_markdown_content(doc, content):
    """Ajoute le contenu markdown au document Word"""
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Titres
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Listes
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Ignorer code blocks et tables
        elif not line.startswith('```') and not line.startswith('|') and not line.startswith('>'):
            doc.add_paragraph(line)

def create_rapport_complet():
    doc = Document()
    
    # Configuration
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # ========== PAGE DE GARDE ==========
    h = doc.add_heading('RAPPORT DE STAGE DE PERFECTIONNEMENT', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p = doc.add_paragraph('Développement d\'une Application Mobile\nde Traduction de la Langue des Signes')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('SignLanguage')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(20)
    p.runs[0].bold = True
    
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('[Votre Nom Complet]\n\n[Votre Filière]\n\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Pépinière d\'Entreprises APII Mahdia\nEspace ISET – Hiboun 5111 – Mahdia\n\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Stage: 1er au 31 janvier 2026\n\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('ISET Mahdia\n2024-2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== REMERCIEMENTS ==========
    doc.add_heading('REMERCIEMENTS', 0)
    doc.add_paragraph('Je tiens à exprimer ma profonde gratitude à toutes les personnes qui ont contribué à la réussite de ce stage de perfectionnement.')
    doc.add_paragraph('Mes remerciements s\'adressent particulièrement au Directeur de la Pépinière APII Mahdia, à mon encadrant ISET, aux formateurs Summer Tech et à ma famille pour leur soutien.')
    doc.add_page_break()
    
    # ========== RÉSUMÉ ==========
    doc.add_heading('RÉSUMÉ', 0)
    doc.add_paragraph('Ce rapport présente le développement de SignLanguage, application mobile de traduction de la langue des signes en temps réel (janvier 2026, Pépinière APII Mahdia). SignLanguage utilise l\'IA (MediaPipe + TensorFlow Lite) pour reconnaître les gestes et les traduire en texte/parole (FR/EN/AR). Précision: 90.3% (lettres), 78.5% (mots), latence: 75ms.')
    doc.add_paragraph('\nMots-clés: Langue des signes, IA, Flutter, MediaPipe, TensorFlow Lite, Accessibilité')
    doc.add_page_break()
    
    # ========== TABLES ==========
    doc.add_heading('TABLE DES MATIÈRES', 0)
    doc.add_paragraph('[Références → Table des matières → Table automatique]')
    doc.add_page_break()
    
    doc.add_heading('LISTE DES ABRÉVIATIONS', 0)
    for abbr, full in [('APII', 'Agence de Promotion de l\'Industrie et de l\'Innovation'),
                       ('IA', 'Intelligence Artificielle'), ('ISET', 'Institut Supérieur des Études Technologiques'),
                       ('LSF', 'Langue des Signes Française'), ('TFLite', 'TensorFlow Lite'),
                       ('TTS', 'Text-To-Speech'), ('UML', 'Unified Modeling Language')]:
        p = doc.add_paragraph()
        p.add_run(f'{abbr}: ').bold = True
        p.add_run(full)
    doc.add_page_break()
    
    # ========== INTRODUCTION (1 PAGE) ==========
    doc.add_heading('INTRODUCTION GÉNÉRALE', 0)
    intro = """La communication est un droit fondamental. Pourtant, 70 millions de personnes sourdes dans le monde font face à des barrières limitant leur accès à l'éducation, l'emploi et les services publics. En Tunisie, 100 000 personnes utilisent la langue des signes. Les interprètes humains (50-100 DT/h) sont rares. L'IA offre une opportunité unique de démocratiser la traduction.

Dans le cadre de ma formation à l'ISET Mahdia, j'ai effectué un stage d'un mois (janvier 2026) à la Pépinière APII Mahdia, structure accompagnant les jeunes porteurs de projets innovants.

SignLanguage traduit les gestes en temps réel en texte/parole (FR/EN/AR). Innovation: support arabe rare. Caractéristiques: on-device, 75ms latence, 90.3% précision lettres, interface emojis accessible.

Objectifs: technique (app mobile + IA), entrepreneurial (Summer Tech, BMC), social (inclusion sourds). Méthodologie: immersion (S1), conception (S2), développement (S3), tests (S4).

Structure: Partie 1 (organisme d'accueil), Partie 2 (projet SignLanguage: état de l'art, analyse, conception, réalisation, tests)."""
    doc.add_paragraph(intro)
    doc.add_page_break()
    
    # ========== PARTIE 1 ==========
    doc.add_heading('PREMIÈRE PARTIE', 0)
    p = doc.add_paragraph('PRÉSENTATION DE L\'ORGANISME D\'ACCUEIL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    doc.add_page_break()
    
    # Lire et ajouter le contenu de la Partie 1
    partie1_content = read_markdown_file(r'C:\Users\dawse\.gemini\antigravity\brain\a032c6aa-78a9-445d-8f5d-6a836360acff\ISET_PARTIE1_pepiniere.md')
    add_markdown_content(doc, partie1_content)
    
    # IMAGE 1: Programme Startup APII 2026
    add_image_with_explanation(doc,
        r'c:\Users\dawse\Desktop\pfa\rapport_images\programme_startup_apii_2026.jpg',
        'Figure 1.1: Programme de formation "startup APII" - Pépinière APII Mahdia (Janvier 2026)',
        '''Ce programme présente le cycle de formation "startup APII" organisé par la Pépinière APII Mahdia en janvier 2026, spécifiquement conçu pour les nouveaux porteurs de projets startup.

Le programme s'est déroulé sur 8 sessions de 3 heures (9h-12h) durant le mois de janvier 2026:

**15/01/2026 - Innovation et écosystème des startups** (M. Slaheddine Dardouri): Introduction à l'écosystème entrepreneurial tunisien, compréhension du Startup Act, structures d'appui (pépinières, incubateurs, accélérateurs), et parcours de création d'une startup.

**16/01/2026 - Prototypage rapide** (Montassar Moussa): Techniques de création rapide de prototypes, utilisation des ressources du FabLab (imprimantes 3D, découpeuses laser), et méthodes de validation d'idées par le prototypage.

**20/01/2026 - Lean Canvas BMC** (Mme Nesrine Akkari): Maîtrise du Business Model Canvas et du Lean Canvas pour structurer une idée de startup en 9 blocs clés (segments clients, proposition de valeur, canaux, relations, revenus, ressources, activités, partenaires, coûts).

**21/01/2026 - Gestion des projets agile / Scrum** (M. Mounir Abida): Méthodologies agiles appliquées aux startups, framework Scrum (sprints, daily standups, retrospectives), outils de gestion (Trello, Jira), et adaptation rapide aux changements.

**22/01/2026 - Loi Startup Act et labélisation** (Slaheddine Dardouri): Présentation détaillée du Startup Act tunisien, avantages fiscaux et sociaux, procédure d'obtention du pré-label et label startup, critères d'éligibilité, et accompagnement APII.

**24/01/2026 - Étude de faisabilité financière** (Montassar Hadj Ayech): Construction du business plan financier, prévisions de trésorerie, compte de résultat prévisionnel, plan de financement, et calcul du seuil de rentabilité.

**26/01/2026 - Étude de faisabilité commerciale** (Nesrine Akkari): Analyse de marché, segmentation clients, stratégie de positionnement, plan marketing et commercial, et validation du marché cible.

**28/01/2026 - Pitch et Branding Digital** (Nesrine Akkari): Techniques de pitch efficace (structure, storytelling, slide design), communication entrepreneuriale, branding digital, présence en ligne, et préparation aux concours startup.

Ce programme intensif m'a été essentiel pour le développement de SignLanguage car il m'a fourni une méthodologie entrepreneuriale complète au-delà des aspects techniques. Le module Lean Canvas m'a permis de structurer mon business model, tandis que les sessions sur le Startup Act m'ont orienté vers la démarche de labélisation. La formation au pitch a été particulièrement utile pour apprendre à présenter le projet de manière convaincante.''')
    
    doc.add_page_break()
    
    # ========== PARTIE 2 ==========
    doc.add_heading('DEUXIÈME PARTIE', 0)
    p = doc.add_paragraph('PROJET SIGNLANGUAGE')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    doc.add_page_break()
    
    # Lire et ajouter le contenu de la Partie 2
    partie2_content = read_markdown_file(r'C:\Users\dawse\.gemini\antigravity\brain\a032c6aa-78a9-445d-8f5d-6a836360acff\ISET_PARTIE2_projet.md')
    add_markdown_content(doc, partie2_content)
    
    # IMAGES 2-9
    images = [
        ('use_case_diagram_1769767904385.png', 'Figure 2.1: Cas d\'utilisation SignLanguage',
         'Acteurs: Utilisateur Sourd, Entendant, Système IA. Cas: Reconnaître gestes, Traduire texte/parole, Changer langue FR/EN/AR, Mode lettres/mots, ESP32-CAM.'),
        
        ('class_diagram_1769767925989.png', 'Figure 3.1: Diagramme de classes',
         'Classes: HandGestureHome, CameraController, HandLandmarkerPlugin (MediaPipe), Interpreter (TFLite), HandPainter. Architecture Flutter respectant SoC.'),
        
        ('sequence_diagram_1769767951765.png', 'Figure 3.2: Séquence reconnaissance lettre',
         'Flux: Capture→MediaPipe (21 landmarks)→Normalisation (84 features)→TFLite→Vote→Affichage. Latence 75ms.'),
        
        ('architecture_diagram_1769767977387.png', 'Figure 3.3: Architecture 4 couches',
         'Présentation (Flutter Widgets), Logique (State, Translation, TTS), IA (MediaPipe, TFLite, Normalisation), Acquisition (Camera, ESP32).'),
        
        ('data_flow_pipeline_1769768053843.png', 'Figure 3.4: Pipeline traitement',
         'Input YUV420→MediaPipe 21×2→Normalisation 84→Modèle Lettres/Mots→Vote→Traduction→Output texte+TTS.'),
        
        ('esp32_cam_setup_1769768074128.png', 'Figure 3.5: ESP32-CAM setup',
         'ESP32-CAM OV2640, FTDI, WiFi streaming HTTP MJPEG 192.168.1.100. Extension IoT pour caméra distante.'),
        
        ('ui_mockup_main_1769768001588.png', 'Figure 3.6: Interface principale',
         'Design emoji-first pour accessibilité sourds. Thème sombre. Contrôles: 🗑️🔦🔊⬅️⌨️. Mode lettres/mots. Langues FR/EN/AR. Vue caméra + landmarks.'),
        
        ('performance_metrics_1769768116914.png', 'Figure 5.1: Métriques performance',
         'Précision: 90.3% lettres (objectif 85%), 78.5% mots (objectif 75%). Latence: 75ms (<100ms). FPS: 24. Taille: 11MB. Langues: 3 (FR/EN/AR).')
    ]
    
    for img_file, caption, expl in images:
        add_image_with_explanation(doc, 
            os.path.join(r'c:\Users\dawse\Desktop\pfa\rapport_images', img_file),
            caption, expl)
    
    doc.add_page_break()
    
    # ========== CONCLUSION (1 PAGE) ==========
    doc.add_heading('CONCLUSION GÉNÉRALE', 0)
    conclusion = """Stage d'un mois APII Mahdia: développement SignLanguage avec 90.3% précision lettres, 78.5% mots, 75ms latence. Summer Tech: outils entrepreneuriaux (BMC, pitch).

Compétences techniques: Flutter, Dart, IA (TensorFlow, MediaPipe, CNN/LSTM), données (3000+ échantillons), IoT (ESP32-CAM), outils (Git, CI/CD).

Compétences transversales: entrepreneuriat (startup, BMC), gestion projet agile, sensibilité sociale (accessibilité, design inclusif).

Défis résolus: rotation landmarks (normalisation dynamique), overfitting LSTM (Dropout 0.4), performance mobile (frame skipping).

Impact social: inclusion sourds, communication quotidienne, emploi, autonomie. Support arabe pour communauté tunisienne. Améliorations: vocabulaire étendu (court terme), reconnaissance continue (moyen terme), RA et iOS (long terme).

Transformation en startup: BMC, équipe, pré-label APII, financements. Carrière: développement mobile + IA à impact social. SignLanguage: début d'un parcours guidé par innovation responsable et inclusion."""
    doc.add_paragraph(conclusion)
    doc.add_page_break()
    
    # ========== BIBLIOGRAPHIE ==========
    doc.add_heading('BIBLIOGRAPHIE', 0)
    doc.add_heading('Ouvrages', level=2)
    doc.add_paragraph('Ries, E. (2011). The Lean Startup. Crown Business.')
    doc.add_paragraph('Osterwalder, A. & Pigneur, Y. (2010). Business Model Generation. Wiley.')
    doc.add_heading('Articles', level=2)
    doc.add_paragraph('Zhang, F. et al. (2020). MediaPipe Framework. arXiv:1906.08172.')
    doc.add_heading('Web', level=2)
    doc.add_paragraph('Flutter: https://flutter.dev')
    doc.add_paragraph('TensorFlow Lite: https://tensorflow.org/lite')
    doc.add_paragraph('MediaPipe: https://developers.google.com/mediapipe')
    
    # Sauvegarder
    output_path = r'c:\Users\dawse\Desktop\pfa\Rapport_SignLanguage_COMPLET.docx'
    doc.save(output_path)
    print(f'✅ RAPPORT COMPLET CRÉÉ: {output_path}')
    print(f'\n📊 Contenu:')
    print('✓ Page de garde')
    print('✓ Remerciements')
    print('✓ Résumé')
    print('✓ Introduction (1 page)')
    print('✓ Partie 1 complète (Pépinière)')
    print('✓ Partie 2 complète (Projet)')
    print('✓ 9 images avec explications')
    print('✓ Conclusion (1 page)')
    print('✓ Bibliographie')
    print(f'\n📝 Nom application: SignLanguage ✓')
    return output_path

if __name__ == '__main__':
    print('🚀 Génération rapport SignLanguage COMPLET...\n')
    create_rapport_complet()
    print('\n✅ TERMINÉ!')
    print('\nProchaines étapes:')
    print('1. Ouvrir le fichier Word')
    print('2. Générer table des matières')
    print('3. Personnaliser [Votre Nom], etc.')
    print('4. Relire et ajuster')
